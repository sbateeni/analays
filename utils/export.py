import os
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches
import pandas as pd
import matplotlib.pyplot as plt
import pdfkit
from io import BytesIO

class AnalysisExporter:
    """فئة لتصدير نتائج التحليل بتنسيقات مختلفة"""
    
    def __init__(self, case, analyses):
        self.case = case
        self.analyses = analyses
        self.export_dir = "exports"
        if not os.path.exists(self.export_dir):
            os.makedirs(self.export_dir)
    
    def export_pdf(self):
        """تصدير التحليل كملف PDF"""
        filename = f"{self.export_dir}/analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        doc = SimpleDocTemplate(filename, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        
        # إنشاء الأنماط
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Arabic', fontName='Arial', alignment=2))  # محاذاة لليمين
        
        # إنشاء المحتوى
        story = []
        
        # إضافة عنوان القضية
        story.append(Paragraph(f"تحليل قضية: {self.case.title}", styles['Arabic']))
        story.append(Spacer(1, 12))
        
        # إضافة معلومات القضية
        story.append(Paragraph(f"رقم القضية: {self.case.case_number}", styles['Arabic']))
        story.append(Paragraph(f"نوع القضية: {self.case.case_type}", styles['Arabic']))
        story.append(Spacer(1, 12))
        
        # إضافة نتائج التحليل
        for analysis in self.analyses:
            story.append(Paragraph(f"المرحلة {analysis.stage}:", styles['Arabic']))
            story.append(Paragraph(analysis.content, styles['Arabic']))
            if analysis.sources:
                story.append(Paragraph("المصادر:", styles['Arabic']))
                for source in analysis.sources:
                    story.append(Paragraph(f"- {source}", styles['Arabic']))
            story.append(Spacer(1, 12))
        
        # بناء الملف
        doc.build(story)
        return filename
    
    def export_word(self):
        """تصدير التحليل كملف Word"""
        filename = f"{self.export_dir}/analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc = Document()
        
        # إضافة عنوان القضية
        doc.add_heading(f"تحليل قضية: {self.case.title}", 0)
        
        # إضافة معلومات القضية
        doc.add_paragraph(f"رقم القضية: {self.case.case_number}")
        doc.add_paragraph(f"نوع القضية: {self.case.case_type}")
        
        # إضافة نتائج التحليل
        for analysis in self.analyses:
            doc.add_heading(f"المرحلة {analysis.stage}:", level=1)
            doc.add_paragraph(analysis.content)
            if analysis.sources:
                doc.add_heading("المصادر:", level=2)
                for source in analysis.sources:
                    doc.add_paragraph(f"- {source}")
        
        # حفظ الملف
        doc.save(filename)
        return filename
    
    def export_excel(self):
        """تصدير التحليل كملف Excel"""
        filename = f"{self.export_dir}/analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        # تحويل البيانات إلى DataFrame
        data = []
        for analysis in self.analyses:
            data.append({
                'المرحلة': analysis.stage,
                'النموذج المستخدم': analysis.model_used,
                'درجة الثقة': analysis.confidence_score,
                'وقت التنفيذ': analysis.execution_time,
                'حالة التحقق': analysis.verification_status,
                'المحتوى': analysis.content
            })
        
        df = pd.DataFrame(data)
        
        # حفظ إلى Excel
        df.to_excel(filename, index=False, sheet_name='التحليل')
        return filename
    
    def create_charts(self):
        """إنشاء رسوم بيانية للتحليل"""
        charts = []
        
        # رسم بياني لدرجات الثقة
        plt.figure(figsize=(10, 6))
        stages = [a.stage for a in self.analyses]
        confidence_scores = [a.confidence_score for a in self.analyses]
        plt.bar(stages, confidence_scores)
        plt.title('درجات الثقة حسب المرحلة')
        plt.xlabel('المرحلة')
        plt.ylabel('درجة الثقة')
        
        # حفظ الرسم البياني
        chart_file = BytesIO()
        plt.savefig(chart_file, format='png', bbox_inches='tight')
        charts.append(chart_file)
        
        # رسم بياني لوقت التنفيذ
        plt.figure(figsize=(10, 6))
        execution_times = [a.execution_time for a in self.analyses]
        plt.plot(stages, execution_times, marker='o')
        plt.title('وقت التنفيذ حسب المرحلة')
        plt.xlabel('المرحلة')
        plt.ylabel('الوقت (ثواني)')
        
        # حفظ الرسم البياني
        chart_file = BytesIO()
        plt.savefig(chart_file, format='png', bbox_inches='tight')
        charts.append(chart_file)
        
        return charts 